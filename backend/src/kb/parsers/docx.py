"""DOCX parser — python-docx.

Reads paragraphs and table cells in document order. Doesn't preserve heading
levels separately (the text content already has them inline).
"""
from __future__ import annotations

import io

from docx import Document as _DocxDocument


def parse_docx(content: bytes) -> str:
    """Extract paragraph + table text from .docx bytes."""
    doc = _DocxDocument(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # Tables: render each row as " | " joined cells, one row per line.
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    return "\n\n".join(parts)
